from docx import Document #pip install python-docx/conda install -c conda-forge python-docx
import datetime
import win32com.client #conda install 

import re
import json
import time
import os
import sys

def saveasdocx(file_path):
     word = win32com.client.Dispatch('Word.Application')
     word.Visible = False
     docx_file = '{0}{1}'.format(file_path, 'x')
     try:
          wordDoc = word.Documents.Open(file_path, False, False, False)
          wordDoc.SaveAs2(docx_file, FileFormat = 16)
          wordDoc.Close()
     except Exception as e: 
          print('Failed to Convert: {0}'.format(file_path))
          print(e)

def getinfo(doc_path):     
#Lấy thông tin của Doucument ở doc_path
     #trả về 4 thông tin: Trích yếu;  Người ký;  Nơi nhận
     #1.Nơi nhận
     trichyeu=""
     noinhan_ct='TH' #Nơi nhận chủ trì
     noinhan_ph=[] #nhận để phối hợp
     noinhan_db=[] #nhận để biết
     noinhan_tn=[] #nhận trong ngành
     nguoi_ky=""
     #bắt đầu chương trình
     document = Document(doc_path)
     if len(document.tables)>1:
        table=document.tables[1] #bảng ở cuối danh sách
        cell=table.rows[0].cells[len(table.columns)-1]
        last_p=len(cell.paragraphs)
        tmp=cell.paragraphs[last_p-1].text.strip()
        if tmp=="Phan Sỹ Bách":
            nguoi_ky="BACPS"
        elif tmp=="Lê Đặng Xuân Tân":
            nguoi_ky="TANLDX"

        #tìm nội dung nơi nhận ở ô Nơi nhận:        
        cell=table.rows[0].cells[0]
        for p in cell.paragraphs:
            if "A0" in p.text and "A0" not in noinhan_tn:
                        noinhan_tn.append("A0")

            if ("b/c" or "bc" or "báo cáo" or "biết") in p.text: #Chỉ có xuất hiện chữ báo cáo hay để biết mới cho vô List nơi nhận để biết, còn lại là cho vô nơi nhận để Thực hiện
                if "các phòng" in p.text.lower():
                    noinhan_db= noinhan_db+["TH","KH","TCKT","DD","PT","CN"]
                elif ("bgđ" in p.text.lower()) or ("ban giám đốc" in p.text.lower()) or ("ban gđ" in p.text.lower()):
                    noinhan_db.append("BACHPS","TANLDX")
                else:
                    if "TH" in p.text and "TH" not in noinhan_db:
                        noinhan_db.append("TH")
                    if "KH" in p.text and "KH" not in noinhan_db:
                        noinhan_db.append("KH")
                    if ( ("TCKT" in p.text) or ("TC-KT" in p.text)) and "TCKT" not in noinhan_db:
                        noinhan_db.append("TCKT")
                    if "DD" in p.text and "DD" not in noinhan_db:
                        noinhan_db.append("DD")
                    if "PT" in p.text and "PT" not in noinhan_db:
                        noinhan_db.append("PT")
                    if "CN" in p.text and "CN" not in noinhan_db:
                        noinhan_db.append("CN")
                    if "GĐ" in p.text and "GĐ" not in noinhan_db:
                        noinhan_db.append("GD")
                    if "Tân" in p.text and "TANLDX" not in noinhan_db:
                        noinhan_db.append("TANLDX")
                    if "Bách" in p.text and "Bách" not in noinhan_db:
                        noinhan_db.append("BACHPS") 
            else:
                if "các phòng" in p.text.lower():
                    for ph in ["TH","KH","TCKT","DD","PT","CN"]:
                        if ph not in noinhan_ph:
                            noinhan_ph.append(ph)
                elif ("bgđ" in p.text.lower()) or ("ban giám đốc" in p.text.lower()) or ("ban gđ" in p.text.lower()):
                    noinhan_ph.append("BACHPS","TANLDX")
                else:
                    if "TH" in p.text and "TH" not in noinhan_ph:
                        noinhan_ph.append("TH")
                    if "KH" in p.text and "KH" not in noinhan_ph:
                        noinhan_ph.append("KH")
                    if "TCKT" in p.text and "TCKT" not in noinhan_ph:
                        noinhan_ph.append("TCKT")
                    if "DD" in p.text and "DD" not in noinhan_ph:
                        noinhan_ph.append("DD")
                    if "PT" in p.text and "PT" not in noinhan_ph:
                        noinhan_ph.append("PT")
                    if "CN" in p.text and "CN" not in noinhan_ph:
                        noinhan_ph.append("CN")
                    if "GĐ" in p.text and "GĐ" not in noinhan_ph:
                        noinhan_ph.append("BACHPS")
                    if "Tân" in p.text and "TANLDX" not in noinhan_ph:
                        noinhan_ph.append("TANLDX")
     
     ghi=False
     for p in document.paragraphs:
         if ("QUYẾT ĐỊNH" in p.text):
             ghi=True
         elif "GIÁM ĐỐC" in p.text: #dừng lại khi thấy chữ Giám đốc đối với Quyết định
             ghi=False
             break

         if ghi==True:
             if ("QUYẾT ĐỊNH" in p.text):
                trichyeu=(p.text).lower().capitalize()+" "                
             elif trichyeu.strip()=="Quyết định":
                trichyeu=trichyeu+" "+(p.text).lower()     #Dòng ngày sau quyết định
             else:
                trichyeu=trichyeu+" "+p.text       

     trichyeu = re.sub(r'[/\():?]', ' ', trichyeu)
     trichyeu = re.sub('\s+',' ',trichyeu) #chỉ để lại tối đa 1 ký tự trống
     trichyeu=trichyeu.strip()
     trichyeu
       

    
     #for i in range(0,len(noinhan)): #lưu ý là range thì sẽ chạy đến cận trên -1 nên lấy len luôn
     #    print("Nơi nhận",i,noinhan[i])
     return (doc_path, trichyeu,nguoi_ky, noinhan_ct, noinhan_ph,noinhan_db, noinhan_tn)

def travel_folder(folder_path):
    dsvb=[]
    for f in os.listdir(folder_path):
        if f.endswith('.docx') and f.startswith('QD'):
            vb=getinfo(f)    
            dsvb.append({'path_vb':vb[0], 'trich_yeu':vb[1],'nguoi_ky':vb[2],'nn_ct':vb[3],'nn_ph':vb[4],'nn_db':vb[5], 'nn_tn':vb[6]})    
    print(dsvb)
    with open('ds_vb.json','w', encoding='utf8') as f:
        json.dump(dsvb,f, ensure_ascii=False, indent=4)

# travel_folder(".")
        
def getinfo2(doc_path):     
#bắt đầu chương trình
     document = Document(doc_path)
     ds_vb=[]
     if len(document.tables)>0:
        table=document.tables[0] #bảng ở cuối danh sách
        # tiêu đề
        r0=[]
        for i in range(len(table.rows)):
            vb={}
            for j in range(len(table.columns)):
                if i==0: #hàng đầu tiên là tiêu đềs
                    r0.append(table.rows[i].cells[j].text)
                else:
                    txt =table.rows[i].cells[j].text
                 
                    if (r0[j]=='trich_yeu') or (r0[j]=='path_vb'):
                        vb[r0[j]]=txt
                    else:
                        vb[r0[j]]=list(txt.split(" "))

            if i!=0:
                ds_vb.append(vb)
        with open('ds_vb.json','w', encoding='utf8') as f:
            json.dump(ds_vb,f, ensure_ascii=False, indent=4)
info = getinfo2('ds_vb.docx')
print(info)